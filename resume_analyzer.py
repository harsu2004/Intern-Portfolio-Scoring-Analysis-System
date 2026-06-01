import os
import re
import warnings
import logging

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
    import pytesseract
    IMAGE_OCR_AVAILABLE = True
except ImportError:
    Image = None
    pytesseract = None
    IMAGE_OCR_AVAILABLE = False

try:
    import pypdfium2 as pdfium
    PDF_OCR_AVAILABLE = True
except ImportError:
    pdfium = None
    PDF_OCR_AVAILABLE = False

warnings.filterwarnings("ignore")
logging.getLogger("pdfminer").setLevel(logging.ERROR)

TESSERACT_EXE = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if IMAGE_OCR_AVAILABLE and os.path.exists(TESSERACT_EXE):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_EXE

OCR_CONFIG = r"--oem 3 --psm 6"

ROLE_SKILLS = {
    "Digital Marketing Intern": ["seo", "google analytics", "google ads", "meta ads", "email marketing", "canva", "semrush", "ahrefs", "content strategy", "copywriting", "campaign", "digital marketing"],
    "Graphic Design Intern": ["photoshop", "illustrator", "indesign", "figma", "canva", "graphic design", "typography", "branding", "logo", "visual design", "mockup", "adobe"],
    "Web Development Intern": ["html", "css", "javascript", "react", "node", "flask", "django", "php", "mongodb", "rest api", "git", "next.js", "express", "typescript", "bootstrap", "tailwind", "docker", "full stack"],
    "Data Analytics Intern": ["python", "sql", "excel", "power bi", "tableau", "pandas", "numpy", "matplotlib", "scikit-learn", "data analysis", "data visualization", "statistics", "machine learning", "eda", "data cleaning", "mysql", "dashboard"],
    "HR/Operations Intern": ["recruitment", "talent acquisition", "onboarding", "payroll", "employee engagement", "resume screening", "scheduling", "ms office", "google sheets", "performance review", "attendance management", "hr", "operations"],
    "Content Writing Intern": ["content writing", "copywriting", "blogging", "seo writing", "proofreading", "editing", "wordpress", "research", "creative writing", "storytelling", "grammarly", "ms word"],
    "Back End Developer Intern": ["python", "node", "django", "flask", "express", "php", "sql", "mongodb", "postgresql", "rest api", "graphql", "docker", "git", "authentication", "server", "database"],
    "Sales & Marketing Intern": ["sales", "crm", "lead generation", "cold calling", "negotiation", "market research", "ms excel", "google sheets", "presentation", "client communication", "b2b", "b2c", "salesforce", "marketing"],
    "UI/UX Design Intern": ["figma", "sketch", "adobe xd", "wireframe", "prototyping", "user research", "usability testing", "ui", "ux", "information architecture", "interaction design", "design system"],
    "Front End Development Intern": ["html", "css", "javascript", "react", "vue", "angular", "typescript", "tailwind", "bootstrap", "responsive design", "git", "next.js", "dom", "component", "figma"],
    "Social Media Management Intern": ["instagram", "facebook", "linkedin", "twitter", "tiktok", "content calendar", "canva", "hootsuite", "buffer", "community management", "engagement", "analytics", "reels", "social media"],
    "Video Editing Intern": ["premiere pro", "after effects", "final cut pro", "davinci resolve", "video editing", "color grading", "motion graphics", "capcut", "storytelling", "youtube", "reels", "transitions", "audio editing"],
    "Legal Internship": ["legal research", "contract drafting", "case study", "ms word", "litigation", "due diligence", "legal writing", "compliance", "documentation", "intellectual property", "corporate law", "negotiation", "legal drafting"]
}

ROLES = list(ROLE_SKILLS.keys())
INSTITUTION_TOKENS = ["school", "college", "university", "institute", "academy", "pvt", "private", "ltd", "inc", "campus"]

def extract_text(path):
    text = ""
    try:
        path_str = str(path)
        if path_str.lower().endswith(".pdf") and pdfplumber is not None:
            with pdfplumber.open(path_str) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t: text += t + " "
        elif path_str.lower().endswith(".docx") and Document is not None:
            doc = Document(path_str)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif path_str.lower().endswith((".png", ".jpg", ".jpeg", ".webp")) and IMAGE_OCR_AVAILABLE:
            img = Image.open(path_str).convert("RGB")
            text = pytesseract.image_to_string(img, config=OCR_CONFIG)
    except Exception:
        text = ""

    for char in ["•", "●", "✔", "◆", "▪", "–", "\uf0b7"]:
        text = text.replace(char, " ")
    text = text.lower().strip()

    if PDF_OCR_AVAILABLE and IMAGE_OCR_AVAILABLE and path_str.lower().endswith(".pdf") and (not text or len(text) < 50):
        try:
            pdf = pdfium.PdfDocument(path_str)
            ocr_text_parts = []
            for i in range(len(pdf)):
                page = pdf[i]
                pil_image = page.render(scale=2).to_pil()
                ocr_text_parts.append(pytesseract.image_to_string(pil_image, config=OCR_CONFIG))
            ocr_text = " ".join(ocr_text_parts)
            for char in ["•", "●", "✔", "◆", "▪", "–", "\uf0b7"]:
                ocr_text = ocr_text.replace(char, " ")
            text = ocr_text.lower().strip()
        except:
            pass
    return text

def extract_email(text):
    match = re.findall(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    return match[0] if match else "Not provided"

def extract_phone(text):
    matches = re.findall(r"\+?\d[\d \-\(\)]{8,15}\d", text)
    for match in matches:
        digits = re.sub(r"\D", "", match)
        if len(digits) == 10 or (len(digits) == 12 and digits.startswith("91")):
            return match.strip()
    return "Not provided"

def get_candidate_name_from_text(text: str, fallback_filename: str) -> str:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in lines[:10]:
        low = line.lower()
        if any(x in low for x in ["@", "email", "phone", "contact", "github", "linkedin", "resume"]): continue
        if any(tok in low for tok in INSTITUTION_TOKENS): continue
        m = re.match(r"^(name[:\-]\s*)(.+)$", line, flags=re.IGNORECASE)
        candidate = m.group(2).strip() if m else line
        if not (1 <= len(candidate.split()) <= 5): continue
        return candidate.strip().title()
    
    base = os.path.splitext(os.path.basename(fallback_filename))[0]
    return base.replace("_", " ").replace("-", " ").title()

# ==========================================================================
# STRICT ALIGNED PARSING EVALUATOR ENGINE
# ==========================================================================
def analyze_resume_data(text, target_role=None):
    clean_text = text.lower()
    
    # 1. Skill Matrix Calculation (Exact Word Bound Check to avoid fake partial matching)
    skills_map = {}
    for role, skills in ROLE_SKILLS.items():
        matched_count = 0
        for s in skills:
            # Using standard word-boundaries to ensure full terms are present
            if re.search(r'\b' + re.escape(s) + r'\b', clean_text):
                matched_count += 1
        skills_map[role] = round((matched_count / len(skills)) * 100, 1)

    if not target_role or target_role not in ROLE_SKILLS:
        target_role = sorted(skills_map.items(), key=lambda x: x[1], reverse=True)[0][0]

    skill_match = skills_map[target_role]

    # 2. Strict Profile Completeness Formula (Dynamic validation based on your request)
    completeness = 10.0  # Base standard structural value
    
    if extract_email(clean_text) != "Not provided": completeness += 15
    if extract_phone(clean_text) != "Not provided": completeness += 15
    
    # Education Validation: Requires structural keyword + verification of degrees
    if "education" in clean_text or "academic" in clean_text:
        if any(deg in clean_text for deg in ["b.tech", "btech", "bca", "mca", "bsc", "bcom", "b.e", "degree", "university", "college"]):
            completeness += 30
        else:
            completeness += 10 # Penalize if section exists but lacks detail
            
    # Project & Experience validation: Requires keyword + validation of technical execution depth
    if any(x in clean_text for x in ["project", "experience", "internship", "work"]):
        # Dynamic check: Did they write detailed bullet points or just list keywords?
        if len(clean_text) > 800:
            completeness += 30
        else:
            completeness += 15

    completeness = min(completeness, 100.0)

    # 3. Resume Quality Score
    quality = 20.0
    if re.search(r"(\b\d+%\s+|\$\d+|\d+\s*+hour|\b(reduced|increased|saved|managed|built|optimized)\b)", clean_text):
        quality += 30
    if "github.com/" in clean_text or "linkedin.com/in/" in clean_text or "portfolio" in clean_text:
        quality += 30
    if 1000 <= len(clean_text) <= 3000:
        quality += 20
    quality = min(quality, 100.0)

    # 4. ATS Optimization Score Matrix (Strict Section Mapping)
    ats_score = 30.0
    
    # Check explicitly defined core structural headings
    sections = ["education", "experience", "project", "skills", "contact"]
    sections_found = sum(1 for sec in sections if sec in clean_text)
    ats_score += (sections_found * 10) # 50 Points max
    
    if 1200 <= len(clean_text) <= 2800:
        ats_score += 20
        
    # Crucial Penalty Rule: If Skill Weight is extremely low, ATS tracking capability decreases
    if skill_match == 0:
        ats_score = max(15.0, ats_score - 25.0)

    ats_score = min(ats_score, 100.0)

    # 5. Combined Balanced Total Score
    total_score = round(
        (skill_match * 0.40) + 
        (ats_score * 0.25) + 
        (quality * 0.20) + 
        (completeness * 0.15), 1
    )

    if total_score >= 75: fit_lbl, fit_cls = "Strong Fit", "strong"
    elif total_score >= 50: fit_lbl, fit_cls = "Moderate Match", "moderate"
    else: fit_lbl, fit_cls = "Needs Improvement", "low"

    matched_keywords = [s.title() for s in ROLE_SKILLS[target_role] if re.search(r'\b' + re.escape(s) + r'\b', clean_text)]
    missing_keywords = [s.title() for s in ROLE_SKILLS[target_role] if not re.search(r'\b' + re.escape(s) + r'\b', clean_text)]

    return {
        "best_role": target_role,
        "best_score": int(total_score),
        "skill_match": int(skill_match),
        "quality": int(quality),
        "completeness": int(completeness),
        "ats_score": int(ats_score),
        "fit_label": fit_lbl,
        "fit_class": fit_cls,
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills_matched": matched_keywords if matched_keywords else ["Core Vector Base"],
        "skills_missing": missing_keywords,
        "all_scores": [{"role": r, "score": int((skills_map[r]*0.4)+(ats_score*0.25)+(quality*0.2)+(completeness*0.15)), "color": "#3b82f6"} for r in ROLES]
    }